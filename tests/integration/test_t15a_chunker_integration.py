"""
Integration tests for T15A Text Chunker with document loaders.

Tests the chunker's ability to process outputs from all document loaders.
"""

import pytest
from unittest.mock import Mock, patch
from pathlib import Path
import tempfile
import json

from src.core.service_manager import ServiceManager
from src.tools.base_tool import ToolRequest, ToolResult
from src.tools.phase1.t01_pdf_loader_unified import T01PDFLoaderUnified
from src.tools.phase1.t02_word_loader_unified import T02WordLoaderUnified
from src.tools.phase1.t03_text_loader_unified import T03TextLoaderUnified
from src.tools.phase1.t04_markdown_loader_unified import T04MarkdownLoaderUnified
from src.tools.phase1.t05_csv_loader_unified import T05CSVLoaderUnified
from src.tools.phase1.t06_json_loader_unified import T06JSONLoaderUnified
from src.tools.phase1.t07_html_loader_unified import T07HTMLLoaderUnified
from src.tools.phase1.t15a_text_chunker_unified import T15ATextChunkerUnified


class TestT15AChunkerIntegration:
    """Test suite for T15A Text Chunker integration with document loaders"""
    
    def setup_method(self):
        """Set up test fixtures"""
        # Mock service manager
        self.service_manager = Mock()
        
        # Mock services
        self.mock_identity = Mock()
        self.mock_provenance = Mock()
        self.mock_quality = Mock()
        
        self.service_manager.identity_service = self.mock_identity
        self.service_manager.provenance_service = self.mock_provenance
        self.service_manager.quality_service = self.mock_quality
        
        # Configure mock returns
        self.mock_identity.create_mention.return_value = {"mention_id": "m123", "status": "success"}
        self.mock_identity.merge_entities.return_value = {"entity_id": "e123", "status": "success"}
        self.mock_provenance.start_operation.return_value = "op123"
        self.mock_provenance.complete_operation.return_value = {"status": "success"}
        self.mock_quality.assess_confidence.return_value = {
            "status": "success",
            "confidence": 0.85,
            "quality_tier": "HIGH"
        }
        self.mock_quality.propagate_confidence.return_value = 0.8
        
        # Initialize tools
        self.text_chunker = T15ATextChunkerUnified(self.service_manager)
        self.text_loader = T03TextLoaderUnified(self.service_manager)
        self.markdown_loader = T04MarkdownLoaderUnified(self.service_manager)
        self.csv_loader = T05CSVLoaderUnified(self.service_manager)
        self.json_loader = T06JSONLoaderUnified(self.service_manager)
        self.html_loader = T07HTMLLoaderUnified(self.service_manager)
    
    def test_chunk_text_loader_output(self):
        """Test chunking plain text from T03 Text Loader"""
        # Create test file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
            test_text = "This is a test document with enough content. " * 50
            f.write(test_text)
            temp_path = f.name
        
        try:
            # Load text
            load_request = ToolRequest(
                tool_id="T03",
                operation="load",
                input_data={"file_path": temp_path}
            )
            load_result = self.text_loader.execute(load_request)
            
            assert load_result.status == "success"
            document_data = load_result.data["document"]
            document_ref = document_data["document_ref"]
            text_content = document_data["text"]
            
            # Chunk the loaded text
            chunk_request = ToolRequest(
                tool_id="T15A",
                operation="chunk",
                input_data={
                    "text": text_content,
                    "document_ref": document_ref,
                    "document_confidence": document_data.get("confidence", 0.8)
                },
                parameters={"chunk_size": 20, "overlap_size": 5}
            )
            chunk_result = self.text_chunker.execute(chunk_request)
            
            assert chunk_result.status == "success"
            assert chunk_result.data["total_chunks"] > 1
            assert all(chunk["source_document"] == document_ref for chunk in chunk_result.data["chunks"])
            
            # Verify chunk overlap
            chunks = chunk_result.data["chunks"]
            for i in range(1, len(chunks)):
                assert chunks[i]["chunk_index"] == i
                # Verify some overlap exists
                prev_words = chunks[i-1]["text"].split()[-5:]
                curr_words = chunks[i]["text"].split()[:5]
                # At least one word should overlap
                assert any(word in curr_words for word in prev_words)
        
        finally:
            Path(temp_path).unlink()
    
    def test_chunk_markdown_loader_output(self):
        """Test chunking markdown content from T04 Markdown Loader"""
        # Create test markdown file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
            markdown_content = """---
title: Test Document
date: 2025-01-01
---

# Introduction

This is the introduction section with some content.

## Section 1

""" + "More content here. " * 50 + """

## Section 2

""" + "Even more content. " * 50
            
            f.write(markdown_content)
            temp_path = f.name
        
        try:
            # Load markdown
            load_request = ToolRequest(
                tool_id="T04",
                operation="load",
                input_data={"file_path": temp_path}
            )
            load_result = self.markdown_loader.execute(load_request)
            
            assert load_result.status == "success"
            document_data = load_result.data["document"]
            document_ref = document_data["document_ref"]
            text_content = document_data["text"]
            
            # Chunk the markdown text
            chunk_request = ToolRequest(
                tool_id="T15A",
                operation="chunk",
                input_data={
                    "text": text_content,
                    "document_ref": document_ref,
                    "document_confidence": document_data.get("confidence", 0.8)
                },
                parameters={"chunk_size": 30, "overlap_size": 10}
            )
            chunk_result = self.text_chunker.execute(chunk_request)
            
            assert chunk_result.status == "success"
            assert chunk_result.data["total_chunks"] > 1
            
            # Verify markdown structure is preserved in chunks
            full_text = " ".join(chunk["text"] for chunk in chunk_result.data["chunks"])
            assert "# Introduction" in full_text
            assert "## Section 1" in full_text
            assert "## Section 2" in full_text
        
        finally:
            Path(temp_path).unlink()
    
    def test_chunk_csv_loader_output(self):
        """Test chunking CSV data converted to text from T05 CSV Loader"""
        # Create test CSV file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.csv', delete=False) as f:
            # Create CSV with enough data to require chunking
            f.write("Name,Description,Category\n")
            for i in range(50):
                f.write(f"Item{i},This is a long description for item {i} with additional details,Category{i%5}\n")
            temp_path = f.name
        
        try:
            # Load CSV
            load_request = ToolRequest(
                tool_id="T05",
                operation="load",
                input_data={"file_path": temp_path}
            )
            load_result = self.csv_loader.execute(load_request)
            
            assert load_result.status == "success"
            dataset_data = load_result.data["dataset"]
            document_ref = dataset_data["dataset_ref"]
            
            # Convert CSV data to text for chunking
            # CSV loader returns data as structured data, need to convert to text
            csv_data = dataset_data["data"]
            text_lines = []
            
            # Add headers
            if csv_data:
                headers = list(csv_data[0].keys())
                text_lines.append(" | ".join(headers))
                
                # Add data rows
                for row in csv_data:
                    text_lines.append(" | ".join(str(row.get(h, "")) for h in headers))
            
            text_content = "\n".join(text_lines)
            
            # Chunk the CSV text
            chunk_request = ToolRequest(
                tool_id="T15A",
                operation="chunk",
                input_data={
                    "text": text_content,
                    "document_ref": document_ref,
                    "document_confidence": dataset_data.get("confidence", 0.8)
                },
                parameters={"chunk_size": 50, "overlap_size": 10}
            )
            chunk_result = self.text_chunker.execute(chunk_request)
            
            assert chunk_result.status == "success"
            assert chunk_result.data["total_chunks"] > 1
            
            # Verify CSV structure is maintained
            first_chunk = chunk_result.data["chunks"][0]["text"]
            assert "Name" in first_chunk  # Header should be in first chunk
            assert "Description" in first_chunk
        
        finally:
            Path(temp_path).unlink()
    
    def test_chunk_json_loader_output(self):
        """Test chunking JSON data converted to text from T06 JSON Loader"""
        # Create test JSON file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False) as f:
            json_data = {
                "title": "Test Document",
                "sections": [
                    {
                        "name": "Introduction",
                        "content": "This is the introduction. " * 30
                    },
                    {
                        "name": "Main Content",
                        "content": "This is the main content section. " * 50
                    },
                    {
                        "name": "Conclusion",
                        "content": "This is the conclusion. " * 20
                    }
                ]
            }
            json.dump(json_data, f)
            temp_path = f.name
        
        try:
            # Load JSON
            load_request = ToolRequest(
                tool_id="T06",
                operation="load",
                input_data={"file_path": temp_path}
            )
            load_result = self.json_loader.execute(load_request)
            
            assert load_result.status == "success"
            document_data = load_result.data["document"]
            document_ref = document_data["document_ref"]
            
            # Convert JSON to text representation
            json_text = json.dumps(document_data["data"], indent=2)
            
            # Chunk the JSON text
            chunk_request = ToolRequest(
                tool_id="T15A",
                operation="chunk",
                input_data={
                    "text": json_text,
                    "document_ref": document_ref,
                    "document_confidence": document_data.get("confidence", 0.8)
                },
                parameters={"chunk_size": 40, "overlap_size": 10}
            )
            chunk_result = self.text_chunker.execute(chunk_request)
            
            assert chunk_result.status == "success"
            assert chunk_result.data["total_chunks"] > 1
            
            # Verify JSON structure elements appear in chunks
            all_chunks_text = " ".join(chunk["text"] for chunk in chunk_result.data["chunks"])
            assert "Introduction" in all_chunks_text
            assert "Main Content" in all_chunks_text
            assert "Conclusion" in all_chunks_text
        
        finally:
            Path(temp_path).unlink()
    
    def test_chunk_html_loader_output(self):
        """Test chunking HTML content from T07 HTML Loader"""
        # Create test HTML file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False) as f:
            html_content = """<!DOCTYPE html>
<html>
<head>
    <title>Test Document</title>
    <meta name="description" content="Test HTML document">
</head>
<body>
    <h1>Main Title</h1>
    <p>""" + "This is paragraph 1 with lots of content. " * 30 + """</p>
    
    <h2>Section 1</h2>
    <p>""" + "This is section 1 content. " * 40 + """</p>
    
    <h2>Section 2</h2>
    <p>""" + "This is section 2 content. " * 40 + """</p>
</body>
</html>"""
            f.write(html_content)
            temp_path = f.name
        
        try:
            # Load HTML
            load_request = ToolRequest(
                tool_id="T07",
                operation="load",
                input_data={"file_path": temp_path}
            )
            load_result = self.html_loader.execute(load_request)
            
            assert load_result.status == "success"
            document_data = load_result.data["document"]
            document_ref = document_data["document_ref"]
            text_content = document_data["text"]
            
            # Chunk the HTML text
            chunk_request = ToolRequest(
                tool_id="T15A",
                operation="chunk",
                input_data={
                    "text": text_content,
                    "document_ref": document_ref,
                    "document_confidence": document_data.get("confidence", 0.8)
                },
                parameters={"chunk_size": 35, "overlap_size": 8}
            )
            chunk_result = self.text_chunker.execute(chunk_request)
            
            assert chunk_result.status == "success"
            assert chunk_result.data["total_chunks"] > 1
            
            # Verify HTML content is preserved
            all_chunks = " ".join(chunk["text"] for chunk in chunk_result.data["chunks"])
            assert "Main Title" in all_chunks
            assert "Section 1" in all_chunks
            assert "Section 2" in all_chunks
        
        finally:
            Path(temp_path).unlink()
    
    def test_chunk_pdf_loader_output(self):
        """Test chunking PDF content from T01 PDF Loader"""
        # Use real test PDF
        test_pdf_path = "/home/brian/projects/Digimons/tools/examples/pdfs/test_document.pdf"
        
        # Initialize PDF loader
        pdf_loader = T01PDFLoaderUnified(self.service_manager)
        
        # Load PDF
        load_request = ToolRequest(
            tool_id="T01",
            operation="load",
            input_data={"file_path": test_pdf_path}
        )
        load_result = pdf_loader.execute(load_request)
        
        assert load_result.status == "success"
        document_data = load_result.data["document"]
        document_ref = document_data["document_ref"]
        text_content = document_data["text"]
        
        # Chunk the PDF text
        chunk_request = ToolRequest(
            tool_id="T15A",
            operation="chunk",
            input_data={
                "text": text_content,
                "document_ref": document_ref,
                "document_confidence": document_data.get("confidence", 0.8)
            },
            parameters={"chunk_size": 40, "overlap_size": 10}
        )
        chunk_result = self.text_chunker.execute(chunk_request)
        
        assert chunk_result.status == "success"
        assert chunk_result.data["total_chunks"] > 0
        
        # Verify content appears in chunks
        all_chunks = " ".join(chunk["text"] for chunk in chunk_result.data["chunks"])
        assert len(all_chunks) > 0
    
    def test_chunk_word_loader_output(self):
        """Test chunking Word document content from T02 Word Loader"""
        # Create a real Word document for testing
        from docx import Document
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            temp_path = f.name
        
        try:
            # Create Word document
            doc = Document()
            doc.add_heading('Test Document', 0)
            doc.add_paragraph("This is the first paragraph with some content. " * 40)
            doc.add_heading('Section 2', level=1)
            doc.add_paragraph("This is the second paragraph with more content. " * 40)
            doc.save(temp_path)
            
            # Initialize Word loader
            word_loader = T02WordLoaderUnified(self.service_manager)
            
            # Load Word document
            load_request = ToolRequest(
                tool_id="T02",
                operation="load",
                input_data={"file_path": temp_path}
            )
            load_result = word_loader.execute(load_request)
            
            assert load_result.status == "success"
            document_data = load_result.data["document"]
            document_ref = document_data["document_ref"]
            text_content = document_data["text"]
            
            # Chunk the Word text
            chunk_request = ToolRequest(
                tool_id="T15A",
                operation="chunk",
                input_data={
                    "text": text_content,
                    "document_ref": document_ref,
                    "document_confidence": document_data.get("confidence", 0.8)
                },
                parameters={"chunk_size": 30, "overlap_size": 8}
            )
            chunk_result = self.text_chunker.execute(chunk_request)
            
            assert chunk_result.status == "success"
            assert chunk_result.data["total_chunks"] > 1
            
            # Verify content from both paragraphs
            all_chunks = " ".join(chunk["text"] for chunk in chunk_result.data["chunks"])
            assert "first paragraph" in all_chunks
            assert "second paragraph" in all_chunks
            
        finally:
            Path(temp_path).unlink()
    
    def test_chunk_size_consistency_across_loaders(self):
        """Test that chunk sizes are consistent regardless of input source"""
        chunk_size = 25
        overlap_size = 5
        
        # Create similar content in different formats
        test_content = "This is test content. " * 100
        
        results = []
        
        # Test with text loader
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
            f.write(test_content)
            temp_path = f.name
        
        try:
            load_request = ToolRequest(
                tool_id="T03",
                operation="load",
                input_data={"file_path": temp_path}
            )
            load_result = self.text_loader.execute(load_request)
            
            document_data = load_result.data["document"]
            chunk_request = ToolRequest(
                tool_id="T15A",
                operation="chunk",
                input_data={
                    "text": document_data["text"],
                    "document_ref": document_data["document_ref"]
                },
                parameters={"chunk_size": chunk_size, "overlap_size": overlap_size}
            )
            chunk_result = self.text_chunker.execute(chunk_request)
            results.append(chunk_result)
        finally:
            Path(temp_path).unlink()
        
        # Test with markdown loader
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
            f.write(f"# Document\n\n{test_content}")
            temp_path = f.name
        
        try:
            load_request = ToolRequest(
                tool_id="T04",
                operation="load",
                input_data={"file_path": temp_path}
            )
            load_result = self.markdown_loader.execute(load_request)
            
            document_data = load_result.data["document"]
            chunk_request = ToolRequest(
                tool_id="T15A",
                operation="chunk",
                input_data={
                    "text": document_data["text"],
                    "document_ref": document_data["document_ref"]
                },
                parameters={"chunk_size": chunk_size, "overlap_size": overlap_size}
            )
            chunk_result = self.text_chunker.execute(chunk_request)
            results.append(chunk_result)
        finally:
            Path(temp_path).unlink()
        
        # Verify chunk counts are similar (within 10%)
        chunk_counts = [r.data["total_chunks"] for r in results]
        assert max(chunk_counts) - min(chunk_counts) <= max(chunk_counts) * 0.1
        
        # Verify chunk sizes are consistent
        for result in results:
            chunks = result.data["chunks"]
            # All chunks except possibly the last should be close to target size
            for chunk in chunks[:-1]:
                assert abs(chunk["token_count"] - chunk_size) <= 5
    
    def test_provenance_tracking_through_pipeline(self):
        """Test that provenance is properly tracked through load->chunk pipeline"""
        # Create test file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
            f.write("Test content for provenance tracking. " * 30)
            temp_path = f.name
        
        try:
            # Load document
            load_request = ToolRequest(
                tool_id="T03",
                operation="load",
                input_data={"file_path": temp_path}
            )
            load_result = self.text_loader.execute(load_request)
            
            assert load_result.status == "success"
            document_data = load_result.data["document"]
            document_ref = document_data["document_ref"]
            
            # Reset provenance mock to track new calls
            self.mock_provenance.start_operation.reset_mock()
            self.mock_provenance.complete_operation.reset_mock()
            
            # Chunk document
            chunk_request = ToolRequest(
                tool_id="T15A",
                operation="chunk",
                input_data={
                    "text": document_data["text"],
                    "document_ref": document_ref
                },
                parameters={"chunk_size": 20}
            )
            chunk_result = self.text_chunker.execute(chunk_request)
            
            assert chunk_result.status == "success"
            
            # Verify provenance tracking
            self.mock_provenance.start_operation.assert_called_once()
            start_call = self.mock_provenance.start_operation.call_args[1]
            assert start_call["tool_id"] == "T15A"
            assert start_call["used"]["document"] == document_ref
            
            self.mock_provenance.complete_operation.assert_called_once()
            complete_call = self.mock_provenance.complete_operation.call_args[1]
            assert complete_call["success"] == True
            assert len(complete_call["outputs"]) == chunk_result.data["total_chunks"]
        
        finally:
            Path(temp_path).unlink()
    
    def test_quality_propagation_through_pipeline(self):
        """Test that quality scores propagate correctly through pipeline"""
        # Create test file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
            f.write("Test content for quality tracking. " * 30)
            temp_path = f.name
        
        try:
            # Load document
            load_request = ToolRequest(
                tool_id="T03",
                operation="load",
                input_data={"file_path": temp_path}
            )
            load_result = self.text_loader.execute(load_request)
            
            document_data = load_result.data["document"]
            document_confidence = document_data.get("confidence", 0.8)
            
            # Chunk document
            chunk_request = ToolRequest(
                tool_id="T15A",
                operation="chunk",
                input_data={
                    "text": document_data["text"],
                    "document_ref": document_data["document_ref"],
                    "document_confidence": document_confidence
                },
                parameters={"chunk_size": 20}
            )
            chunk_result = self.text_chunker.execute(chunk_request)
            
            assert chunk_result.status == "success"
            
            # Verify quality propagation was called
            self.mock_quality.propagate_confidence.assert_called()
            
            # Verify all chunks have confidence scores
            chunks = chunk_result.data["chunks"]
            for chunk in chunks:
                assert "confidence" in chunk
                assert 0 < chunk["confidence"] <= 1.0
                assert "quality_tier" in chunk
        
        finally:
            Path(temp_path).unlink()
    
    def test_error_handling_in_pipeline(self):
        """Test error handling when chunking fails in pipeline"""
        # Create test file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
            f.write("Valid content")
            temp_path = f.name
        
        try:
            # Load document successfully
            load_request = ToolRequest(
                tool_id="T03",
                operation="load",
                input_data={"file_path": temp_path}
            )
            load_result = self.text_loader.execute(load_request)
            
            assert load_result.status == "success"
            
            document_data = load_result.data["document"]
            
            # Try to chunk with empty text (simulating data loss)
            chunk_request = ToolRequest(
                tool_id="T15A",
                operation="chunk",
                input_data={
                    "text": "",  # Empty text
                    "document_ref": document_data["document_ref"]
                }
            )
            chunk_result = self.text_chunker.execute(chunk_request)
            
            assert chunk_result.status == "error"
            assert chunk_result.error_code == "EMPTY_TEXT"
        
        finally:
            Path(temp_path).unlink()
    
    def test_performance_with_large_documents(self):
        """Test performance when processing large documents through pipeline"""
        import time
        
        # Create large test file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
            # Create ~100KB of text
            large_content = "This is test content. " * 5000
            f.write(large_content)
            temp_path = f.name
        
        try:
            # Load document
            start_time = time.time()
            load_request = ToolRequest(
                tool_id="T03",
                operation="load",
                input_data={"file_path": temp_path}
            )
            load_result = self.text_loader.execute(load_request)
            load_time = time.time() - start_time
            
            assert load_result.status == "success"
            assert load_time < 2.0  # Should load within 2 seconds
            
            document_data = load_result.data["document"]
            
            # Chunk document
            start_time = time.time()
            chunk_request = ToolRequest(
                tool_id="T15A",
                operation="chunk",
                input_data={
                    "text": document_data["text"],
                    "document_ref": document_data["document_ref"]
                },
                parameters={"chunk_size": 100, "overlap_size": 20}
            )
            chunk_result = self.text_chunker.execute(chunk_request)
            chunk_time = time.time() - start_time
            
            assert chunk_result.status == "success"
            assert chunk_time < 5.0  # Should chunk within 5 seconds
            assert chunk_result.data["total_chunks"] > 10  # Should produce many chunks
            
            # Total pipeline time should be reasonable
            assert load_time + chunk_time < 7.0
        
        finally:
            Path(temp_path).unlink()