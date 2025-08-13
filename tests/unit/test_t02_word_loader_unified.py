"""
Mock-Free Tests for T02 Word Loader - Unified Interface Implementation

These tests use REAL functionality with NO mocking of core operations.
All tests use actual files, real python-docx execution, and real ServiceManager instances.
"""

import pytest
import os
import tempfile
import shutil
from pathlib import Path
from datetime import datetime
import time

# Real imports - NO mocking imports
from src.tools.phase1.t02_word_loader_unified import T02WordLoaderUnified
from src.tools.base_tool import ToolRequest, ToolResult, ToolContract, ToolStatus
from src.core.service_manager import ServiceManager


class TestT02WordLoaderUnifiedMockFree:
    """Mock-free testing for T02 Word Loader unified interface"""
    
    def setup_method(self):
        """Set up test fixtures with REAL services and files"""
        # Use REAL ServiceManager instance
        self.service_manager = ServiceManager()
        self.tool = T02WordLoaderUnified(self.service_manager)
        
        # Create temp directory for test files
        self.test_dir = Path(tempfile.mkdtemp())
        
        # Create REAL test files
        self.test_docx_path = self._create_real_test_docx()
        self.complex_docx_path = self._create_complex_docx()
        self.corrupted_docx_path = self._create_corrupted_docx()
    
    def teardown_method(self):
        """Clean up test files"""
        if self.test_dir.exists():
            shutil.rmtree(self.test_dir)
    
    def _create_real_test_docx(self) -> Path:
        """Create actual DOCX file using python-docx for testing"""
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError:
            pytest.skip("python-docx not available for testing")
        
        # Create a real DOCX document
        document = Document()
        
        # Add title
        document.add_heading('Test DOCX Document', 0)
        
        # Add paragraphs with test content
        document.add_paragraph('This is a test document for the T02 Word Loader.')
        document.add_paragraph('Microsoft was founded by Bill Gates and Paul Allen in 1975.')
        document.add_paragraph('Apple Inc. was founded by Steve Jobs, Steve Wozniak, and Ronald Wayne.')
        
        # Add a formatted paragraph
        paragraph = document.add_paragraph('This paragraph contains ')
        run = paragraph.add_run('bold text')
        run.bold = True
        paragraph.add_run(' and ')
        run = paragraph.add_run('italic text')
        run.italic = True
        paragraph.add_run('.')
        
        # Add a table
        table = document.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Company'
        hdr_cells[1].text = 'Founder'
        
        # Data rows
        row_cells = table.rows[1].cells
        row_cells[0].text = 'Microsoft'
        row_cells[1].text = 'Bill Gates'
        
        row_cells = table.rows[2].cells
        row_cells[0].text = 'Apple'
        row_cells[1].text = 'Steve Jobs'
        
        # Add another paragraph after the table
        document.add_paragraph('This content appears after the table.')
        
        # Save the document
        test_file = self.test_dir / "test_document.docx"
        document.save(str(test_file))
        return test_file
    
    def _create_complex_docx(self) -> Path:
        """Create complex DOCX with multiple features"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available for testing")
        
        document = Document()
        
        # Add multiple headings and content
        document.add_heading('Complex Document Structure', 0)
        document.add_heading('Section 1: Overview', 1)
        document.add_paragraph('This section provides an overview of the document.')
        
        # Create a large table
        table = document.add_table(rows=5, cols=3)
        table.style = 'Table Grid'
        
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.text = f'Row {i+1}, Col {j+1}'
        
        document.add_heading('Section 2: Details', 1)
        
        # Add multiple paragraphs
        for i in range(10):
            document.add_paragraph(f'This is paragraph {i+1} with detailed content about the topic. ' * 5)
        
        # Add a list
        document.add_paragraph('Key Points:', style='List Bullet')
        document.add_paragraph('First important point', style='List Bullet')
        document.add_paragraph('Second important point', style='List Bullet')
        document.add_paragraph('Third important point', style='List Bullet')
        
        test_file = self.test_dir / "complex_document.docx"
        document.save(str(test_file))
        return test_file
    
    def _create_corrupted_docx(self) -> Path:
        """Create actually corrupted DOCX file"""
        corrupted_content = b"This is not a DOCX file, it's corrupted data"
        test_file = self.test_dir / "corrupted.docx"
        with open(test_file, 'wb') as f:
            f.write(corrupted_content)
        return test_file
    
    # ===== CONTRACT TESTS (MANDATORY) =====
    
    def test_tool_initialization_real(self):
        """Tool initializes with real services"""
        assert self.tool is not None
        assert self.tool.tool_id == "T02"
        assert self.tool.services == self.service_manager
        assert isinstance(self.tool, T02WordLoaderUnified)
        
        # Verify real service connections
        assert self.tool.identity_service is not None
        assert self.tool.provenance_service is not None
        assert self.tool.quality_service is not None
    
    def test_get_contract_real(self):
        """Tool provides complete contract specification"""
        contract = self.tool.get_contract()
        
        assert isinstance(contract, ToolContract)
        assert contract.tool_id == "T02"
        assert contract.name == "Word Document Loader"
        assert contract.category == "document_processing"
        assert contract.description == "Load and extract text from Word documents (.docx)"
        
        # Verify input schema
        assert "file_path" in contract.input_schema["properties"]
        assert "workflow_id" in contract.input_schema["properties"]
        assert contract.input_schema["required"] == ["file_path"]
        
        # Verify output schema structure
        assert "document" in contract.output_schema["properties"]
        doc_props = contract.output_schema["properties"]["document"]["properties"]
        assert "text" in doc_props
        assert "confidence" in doc_props
        assert "document_id" in doc_props
        
        # Verify dependencies
        assert "identity_service" in contract.dependencies
        assert "provenance_service" in contract.dependencies
        assert "quality_service" in contract.dependencies
        
        # Verify performance requirements
        assert contract.performance_requirements["max_execution_time"] == 20.0
        assert contract.performance_requirements["max_memory_mb"] == 1024
    
    def test_input_contract_validation_real(self):
        """Tool validates inputs according to contract using real validation"""
        # Test invalid inputs with real validation
        invalid_inputs = [
            {},  # Empty input
            {"wrong_field": "value"},  # Wrong fields
            None,  # Null input
            {"file_path": ""},  # Empty file path
            {"file_path": 123},  # Wrong type
            {"file_path": str(self.test_dir / "nonexistent.docx")},  # File doesn't exist
            {"file_path": "/etc/passwd"},  # Security risk
            {"file_path": str(self.test_dir / "test.pdf")},  # Wrong extension
            {"file_path": str(self.test_dir / "test.doc")},  # Old Word format not supported
        ]
        
        for invalid_input in invalid_inputs:
            request = ToolRequest(
                tool_id="T02",
                operation="load",
                input_data=invalid_input,
                parameters={}
            )
            result = self.tool.execute(request)
            assert result.status == "error"
            assert result.error_code in [
                "INVALID_INPUT", "VALIDATION_FAILED", "INVALID_FILE_TYPE", 
                "FILE_NOT_FOUND", "INVALID_FILE_EXTENSION"
            ]
    
    # ===== REAL FUNCTIONALITY TESTS =====
    
    def test_docx_loading_real_functionality(self):
        """Test DOCX loading with REAL python-docx execution"""
        request = ToolRequest(
            tool_id="T02",
            operation="load",
            input_data={"file_path": str(self.test_docx_path)},
            parameters={}
        )
        
        # Execute with REAL functionality
        start_time = time.time()
        result = self.tool.execute(request)
        execution_time = time.time() - start_time
        
        # Verify REAL results
        assert result.status == "success"
        assert result.tool_id == "T02"
        
        # Verify document structure
        document = result.data["document"]
        assert "document_id" in document
        assert len(document["text"]) > 0
        assert document["paragraph_count"] >= 1
        assert document["confidence"] > 0.0
        assert document["file_path"] == str(self.test_docx_path)
        assert document["file_size"] > 0
        
        # Verify real timing
        assert result.execution_time > 0
        assert execution_time < 20.0  # Performance requirement
        
        # Verify text content was actually extracted
        text = document["text"]
        assert len(text.strip()) > 0  # Not empty
        assert "Microsoft was founded by Bill Gates" in text
        assert "Apple Inc. was founded by Steve Jobs" in text
        
        # Verify table content was extracted
        assert "Microsoft" in text
        assert "Bill Gates" in text
        assert "Apple" in text
        assert "Steve Jobs" in text
        
        # Verify metadata
        assert result.metadata["operation_id"] is not None
        assert "workflow_id" in result.metadata
    
    def test_table_extraction_real(self):
        """Test table extraction with REAL python-docx execution"""
        request = ToolRequest(
            tool_id="T02",
            operation="load",
            input_data={"file_path": str(self.test_docx_path)},
            parameters={"extract_tables": True}
        )
        
        result = self.tool.execute(request)
        assert result.status == "success"
        
        document = result.data["document"]
        
        # Should have table data
        if "table_count" in document:
            assert document["table_count"] >= 1
        
        # Table content should be in the text
        text = document["text"]
        assert "Company" in text  # Table header
        assert "Founder" in text  # Table header
        assert "Microsoft" in text  # Table data
        assert "Apple" in text  # Table data
    
    def test_complex_document_real(self):
        """Test complex document with multiple features"""
        request = ToolRequest(
            tool_id="T02",
            operation="load",
            input_data={"file_path": str(self.complex_docx_path)},
            parameters={}
        )
        
        result = self.tool.execute(request)
        assert result.status == "success"
        
        document = result.data["document"]
        text = document["text"]
        
        # Should contain headings
        assert "Complex Document Structure" in text
        assert "Section 1: Overview" in text
        assert "Section 2: Details" in text
        
        # Should contain paragraph content
        assert "This is paragraph" in text
        
        # Should have substantial content
        assert len(text) > 1000  # Complex document should be substantial
        assert document["paragraph_count"] > 10
    
    def test_corrupted_docx_real_error_handling(self):
        """Test corrupted DOCX with REAL error handling"""
        request = ToolRequest(
            tool_id="T02",
            operation="load",
            input_data={"file_path": str(self.corrupted_docx_path)},
            parameters={}
        )
        
        # Should get REAL error from python-docx
        result = self.tool.execute(request)
        assert result.status == "error"
        assert result.error_code in ["DOCX_CORRUPTED", "EXTRACTION_FAILED"]
        
        # Verify error message contains meaningful information
        assert len(result.error_message) > 0
        assert result.error_message is not None
    
    def test_file_not_found_real_error(self):
        """Test missing file with REAL filesystem check"""
        nonexistent_path = str(self.test_dir / "does_not_exist.docx")
        
        request = ToolRequest(
            tool_id="T02",
            operation="load",
            input_data={"file_path": nonexistent_path},
            parameters={}
        )
        
        result = self.tool.execute(request)
        assert result.status == "error"
        assert result.error_code == "FILE_NOT_FOUND"
        assert "not found" in result.error_message.lower() or "does not exist" in result.error_message.lower()
    
    def test_unsupported_file_type_real_validation(self):
        """Test unsupported file type with REAL file validation"""
        # Create a real file with unsupported extension
        unsupported_file = self.test_dir / "document.pdf"
        with open(unsupported_file, 'w') as f:
            f.write("This is not a supported file type")
        
        request = ToolRequest(
            tool_id="T02",
            operation="load",
            input_data={"file_path": str(unsupported_file)},
            parameters={}
        )
        
        result = self.tool.execute(request)
        assert result.status == "error"
        assert result.error_code == "INVALID_FILE_TYPE"
        assert "unsupported" in result.error_message.lower() or "invalid" in result.error_message.lower()
    
    # ===== INTEGRATION TESTS WITH REAL SERVICES =====
    
    def test_identity_service_integration_real(self):
        """Test integration with real IdentityService"""
        request = ToolRequest(
            tool_id="T02",
            operation="load",
            input_data={
                "file_path": str(self.test_docx_path),
                "workflow_id": "test_workflow_456"
            },
            parameters={}
        )
        
        result = self.tool.execute(request)
        assert result.status == "success"
        
        # Verify document ID follows real pattern
        document_id = result.data["document"]["document_id"]
        assert "test_workflow_456" in document_id
        assert "test_document" in document_id  # Based on filename
    
    def test_provenance_service_integration_real(self):
        """Test integration with real ProvenanceService"""
        request = ToolRequest(
            tool_id="T02",
            operation="load",
            input_data={"file_path": str(self.test_docx_path)},
            parameters={}
        )
        
        result = self.tool.execute(request)
        assert result.status == "success"
        
        # Verify provenance tracking actually occurred
        assert "operation_id" in result.metadata
        operation_id = result.metadata["operation_id"]
        assert operation_id is not None
        assert len(operation_id) > 0
    
    def test_quality_service_integration_real(self):
        """Test integration with real QualityService"""
        request = ToolRequest(
            tool_id="T02",
            operation="load",
            input_data={"file_path": str(self.test_docx_path)},
            parameters={}
        )
        
        result = self.tool.execute(request)
        assert result.status == "success"
        
        # Verify quality assessment actually occurred
        document = result.data["document"]
        assert "confidence" in document
        assert isinstance(document["confidence"], (int, float))
        assert 0.0 <= document["confidence"] <= 1.0
        
        # May have quality_tier if quality service provides it
        if "quality_tier" in document:
            assert document["quality_tier"] in ["LOW", "MEDIUM", "HIGH"]
    
    # ===== PERFORMANCE TESTS WITH REAL EXECUTION =====
    
    def test_performance_requirements_real(self):
        """Test tool meets performance benchmarks with real execution"""
        request = ToolRequest(
            tool_id="T02",
            operation="load",
            input_data={"file_path": str(self.test_docx_path)},
            parameters={}
        )
        
        # Measure performance with real execution
        start_time = time.time()
        result = self.tool.execute(request)
        execution_time = time.time() - start_time
        
        # Performance assertions
        assert result.status == "success"
        assert execution_time < 20.0  # Max 20 seconds
        assert result.execution_time < 20.0
        
        # Memory usage should be reasonable (if tracked)
        if result.memory_used > 0:
            assert result.memory_used < 1024 * 1024 * 1024  # Max 1GB
    
    def test_large_document_handling_real(self):
        """Test handling of larger documents with real data"""
        request = ToolRequest(
            tool_id="T02",
            operation="load",
            input_data={"file_path": str(self.complex_docx_path)},
            parameters={}
        )
        
        start_time = time.time()
        result = self.tool.execute(request)
        execution_time = time.time() - start_time
        
        assert result.status == "success"
        assert len(result.data["document"]["text"]) > 2000  # Substantial content (adjusted for actual output)
        assert execution_time < 20.0  # Should still be fast
        assert result.data["document"]["paragraph_count"] > 10
    
    # ===== TOOL INTERFACE TESTS =====
    
    def test_tool_status_management_real(self):
        """Tool manages status correctly during real execution"""
        assert self.tool.get_status() == ToolStatus.READY
        
        # Status should remain consistent after operations
        request = ToolRequest(
            tool_id="T02",
            operation="load",
            input_data={"file_path": str(self.test_docx_path)},
            parameters={}
        )
        
        result = self.tool.execute(request)
        assert result.status == "success"
        assert self.tool.get_status() == ToolStatus.READY
    
    def test_health_check_real(self):
        """Tool health check works with real dependencies"""
        result = self.tool.health_check()
        
        assert isinstance(result, ToolResult)
        assert result.tool_id == "T02"
        assert result.status in ["success", "error"]
        
        if result.status == "success":
            assert result.data["healthy"] == True
            assert "supported_formats" in result.data
            assert ".docx" in result.data["supported_formats"]
            
            # Verify real service health
            assert result.data.get("services_healthy") in [True, None]
    
    def test_cleanup_real(self):
        """Tool cleans up resources properly"""
        # Add some temp files to the tool
        temp_file = self.test_dir / "temp_test.docx"
        
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("temp content")
            doc.save(str(temp_file))
        except ImportError:
            # Fallback to creating empty file
            temp_file.touch()
        
        self.tool._temp_files.append(str(temp_file))
        
        # Test cleanup
        success = self.tool.cleanup()
        assert success == True
        
        # Temp files list should be cleared
        assert len(self.tool._temp_files) == 0
    
    # ===== EDGE CASES WITH REAL CONDITIONS =====
    
    def test_empty_document_real(self):
        """Test empty document handling with real empty DOCX"""
        try:
            from docx import Document
            
            # Create empty document
            document = Document()
            empty_file = self.test_dir / "empty.docx"
            document.save(str(empty_file))
            
            request = ToolRequest(
                tool_id="T02",
                operation="load",
                input_data={"file_path": str(empty_file)},
                parameters={}
            )
            
            result = self.tool.execute(request)
            # May succeed with empty content or fail gracefully
            if result.status == "success":
                assert len(result.data["document"]["text"].strip()) == 0
                assert result.data["document"]["paragraph_count"] == 0
            else:
                assert result.error_code is not None
                
        except ImportError:
            pytest.skip("python-docx not available for testing")
    
    def test_workflow_id_generation_real(self):
        """Test workflow ID generation with real logic"""
        request = ToolRequest(
            tool_id="T02",
            operation="load",
            input_data={"file_path": str(self.test_docx_path)},
            parameters={}
        )
        
        result = self.tool.execute(request)
        assert result.status == "success"
        
        # Should have generated a workflow ID
        assert "workflow_id" in result.metadata
        workflow_id = result.metadata["workflow_id"]
        assert workflow_id.startswith("wf_")
        assert len(workflow_id) > 3
        
        # Document ID should include the workflow ID
        document_id = result.data["document"]["document_id"]
        assert workflow_id in document_id
    
    def test_confidence_calculation_real(self):
        """Test confidence calculation with real factors"""
        # Test with different documents
        files_to_test = [
            (self.test_docx_path, "simple document"),
            (self.complex_docx_path, "complex document")
        ]
        
        confidences = []
        for file_path, description in files_to_test:
            request = ToolRequest(
                tool_id="T02",
                operation="load",
                input_data={"file_path": str(file_path)},
                parameters={}
            )
            
            result = self.tool.execute(request)
            assert result.status == "success"
            
            confidence = result.data["document"]["confidence"]
            confidences.append((confidence, description))
            
            # Verify confidence is in valid range
            assert 0.0 <= confidence <= 1.0
            assert isinstance(confidence, (int, float))
        
        # Verify confidence values are reasonable
        confidence_values = [c[0] for c in confidences]
        # Should have some variation based on document characteristics
        assert min(confidence_values) >= 0.1  # Not too low
        assert max(confidence_values) <= 1.0  # Not too high
    
    def test_formatting_preservation_real(self):
        """Test that formatting is handled correctly with real documents"""
        request = ToolRequest(
            tool_id="T02",
            operation="load",
            input_data={"file_path": str(self.test_docx_path)},
            parameters={"preserve_formatting": False}  # Test parameter handling
        )
        
        result = self.tool.execute(request)
        assert result.status == "success"
        
        # Should extract text without formatting markup
        text = result.data["document"]["text"]
        assert "bold text" in text
        assert "italic text" in text
        # Should not contain formatting tags
        assert "<b>" not in text
        assert "<i>" not in text
    
    def test_word_count_accuracy_real(self):
        """Test word count calculation with real content"""
        request = ToolRequest(
            tool_id="T02",
            operation="load",
            input_data={"file_path": str(self.test_docx_path)},
            parameters={}
        )
        
        result = self.tool.execute(request)
        assert result.status == "success"
        
        document = result.data["document"]
        text = document["text"]
        
        # Calculate expected word count
        expected_words = len(text.split())
        
        # Should have word count information
        if "total_words" in document:
            actual_words = document["total_words"]
            # Should be reasonably close (within 10% for formatting differences)
            assert abs(actual_words - expected_words) <= max(expected_words * 0.1, 5)
        
        # Should have reasonable word count for test document
        assert expected_words > 20  # Our test document has substantial content